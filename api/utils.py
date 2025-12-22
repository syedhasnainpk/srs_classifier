import os
import pickle
import numpy as np
import faiss
from sentence_transformers import SentenceTransformer
import logging
from typing import List, Dict, Tuple
import json

logger = logging.getLogger(__name__)


class FreeVectorStore:
    def __init__(self, embedding_dim=384, persist_dir="./vector_store"):
        self.embedding_dim = embedding_dim
        self.persist_dir = persist_dir
        self.index = None
        self.documents = []  # Store document metadata
        self.embedder = SentenceTransformer('all-MiniLM-L6-v2')

        # Create persist directory
        os.makedirs(persist_dir, exist_ok=True)

        # Load existing index if available
        self.load_index()

    def load_index(self):
        """Load existing FAISS index and documents"""
        index_path = os.path.join(self.persist_dir, "faiss_index.bin")
        docs_path = os.path.join(self.persist_dir, "documents.json")

        if os.path.exists(index_path) and os.path.exists(docs_path):
            try:
                self.index = faiss.read_index(index_path)
                with open(docs_path, 'r') as f:
                    self.documents = json.load(f)
                logger.info(f"✅ Loaded existing index with {len(self.documents)} documents")
            except Exception as e:
                logger.warning(f"Failed to load existing index: {e}")
                self.create_new_index()
        else:
            self.create_new_index()

    def create_new_index(self):
        """Create new FAISS index"""
        self.index = faiss.IndexFlatIP(self.embedding_dim)  # Inner product similarity
        self.documents = []
        logger.info("✅ Created new FAISS index")

    def add_documents(self, texts: List[str], metadata: List[Dict]):
        """Add documents to the vector store"""
        if not texts or len(texts) != len(metadata):
            raise ValueError("Texts and metadata must have same length")

        # Generate embeddings
        embeddings = self.embedder.encode(texts, convert_to_numpy=True, normalize_embeddings=True)

        # Add to FAISS index
        self.index.add(embeddings.astype('float32'))

        # Store metadata
        for i, (text, meta) in enumerate(zip(texts, metadata)):
            doc_id = len(self.documents)
            self.documents.append({
                'id': doc_id,
                'text': text,
                'metadata': meta
            })

        # Persist to disk
        self.save_index()
        logger.info(f"✅ Added {len(texts)} documents to vector store")

    def search(self, query: str, k: int = 5, score_threshold: float = 0.3) -> List[Dict]:
        """Search for similar documents"""
        if self.index.ntotal == 0:
            return []

        # Generate query embedding
        query_embedding = self.embedder.encode([query], convert_to_numpy=True, normalize_embeddings=True)

        # Search
        scores, indices = self.index.search(query_embedding.astype('float32'), k)

        results = []
        for score, idx in zip(scores[0], indices[0]):
            if idx != -1 and score >= score_threshold:  # -1 means no match found
                doc = self.documents[idx]
                results.append({
                    'text': doc['text'],
                    'score': float(score),
                    'metadata': doc['metadata']
                })

        return results

    def save_index(self):
        """Save FAISS index and documents to disk"""
        try:
            index_path = os.path.join(self.persist_dir, "faiss_index.bin")
            docs_path = os.path.join(self.persist_dir, "documents.json")

            faiss.write_index(self.index, index_path)
            with open(docs_path, 'w') as f:
                json.dump(self.documents, f)

            logger.info("✅ Vector store saved to disk")
        except Exception as e:
            logger.error(f"Failed to save vector store: {e}")


# Global vector store instance
vector_store = FreeVectorStore()


# Keep existing utility functions
def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    try:
        from docx import Document
        doc = Document(file_path)
        full_text = []

        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))

        text = "\n".join(full_text)
        if not text.strip():
            raise ValueError("No text content found in DOCX file")

        logger.info(f"Successfully extracted {len(text)} characters from DOCX")
        return text

    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {str(e)}")
        raise Exception(f"DOCX extraction failed: {str(e)}")


def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    try:
        import PyPDF2
        with open(file_path, 'rb') as file:
            try:
                pdf_reader = PyPDF2.PdfReader(file)
                pages = pdf_reader.pages
            except AttributeError:
                pdf_reader = PyPDF2.PdfFileReader(file)
                pages = [pdf_reader.getPage(i) for i in range(pdf_reader.numPages)]

            full_text = []
            for page in pages:
                try:
                    text = page.extract_text() if hasattr(page, 'extract_text') else page.extractText()
                    if text and text.strip():
                        full_text.append(text.strip())
                except Exception as e:
                    logger.warning(f"Failed to extract text from a page: {e}")
                    continue

            if not full_text:
                raise ValueError("No text could be extracted from PDF")

            text = "\n".join(full_text)
            logger.info(f"Successfully extracted {len(text)} characters from PDF")
            return text

    except Exception as e:
        logger.error(f"Failed to extract text from PDF: {str(e)}")
        raise Exception(f"PDF extraction failed: {str(e)}")


def chunk_text(text, max_chunk_size=500, overlap_size=50):
    """Chunk text with proper overlap"""
    import re

    if not text or not text.strip():
        raise ValueError("Empty text provided for chunking")

    # Clean text
    text = re.sub(r'\s+', ' ', text.strip())

    # Split into sentences
    sentences = re.split(r'(?<=[.!?])\s+', text)
    if not sentences:
        sentences = text.split('. ')

    chunks = []
    i = 0

    while i < len(sentences):
        current_chunk = ""
        sentence_count = 0

        # Build chunk
        while i + sentence_count < len(sentences):
            sentence = sentences[i + sentence_count]
            potential_chunk = current_chunk + " " + sentence if current_chunk else sentence

            if len(potential_chunk) <= max_chunk_size:
                current_chunk = potential_chunk
                sentence_count += 1
            else:
                break

        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        # Move forward with overlap
        if sentence_count <= 1:
            i += 1
        else:
            overlap_sentences = max(1, min(sentence_count // 2, overlap_size // 50))
            i += max(1, sentence_count - overlap_sentences)

    if not chunks:
        raise ValueError("No valid chunks could be created")

    logger.info(f"Created {len(chunks)} chunks from text")
    return chunks